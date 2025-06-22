# /backend/app/routers/export_doc.py

from typing import Literal
from uuid import UUID
import io
import re
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, status, Response
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
import sentry_sdk
from html.parser import HTMLParser

from ..auth import get_current_user

# Import file generation libraries
try:
    import docx
    from docx.shared import Inches
except ImportError:
    docx = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
except ImportError:
    SimpleDocTemplate = None

# Content size limits (reasonable limits for file generation)
MAX_CONTENT_LENGTH = 1_000_000  # 1MB of text content
MAX_TITLE_LENGTH = 255

router = APIRouter(prefix="/export", tags=["export"])


class RichTextParser(HTMLParser):
    """HTML parser to extract text with formatting information."""
    
    def __init__(self):
        super().__init__()
        self.reset()
        self.text_parts = []
        self.current_formatting = set()
        
    def handle_starttag(self, tag, attrs):
        if tag in ['strong', 'b']:
            self.current_formatting.add('bold')
        elif tag in ['em', 'i']:
            self.current_formatting.add('italic')
        elif tag == 'u':
            self.current_formatting.add('underline')
        elif tag in ['s', 'strike']:
            self.current_formatting.add('strikethrough')
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.current_formatting.add('heading')
            self.current_formatting.add(f'heading-{tag[1]}')
        elif tag == 'br':
            self.text_parts.append({
                'text': '\n',
                'formatting': self.current_formatting.copy(),
                'type': 'linebreak'
            })
        elif tag == 'p':
            # Start of paragraph
            pass
            
    def handle_endtag(self, tag):
        if tag in ['strong', 'b']:
            self.current_formatting.discard('bold')
        elif tag in ['em', 'i']:
            self.current_formatting.discard('italic')
        elif tag == 'u':
            self.current_formatting.discard('underline')
        elif tag in ['s', 'strike']:
            self.current_formatting.discard('strikethrough')
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.current_formatting.discard('heading')
            self.current_formatting.discard(f'heading-{tag[1]}')
            # Add line break after heading
            self.text_parts.append({
                'text': '\n',
                'formatting': set(),
                'type': 'linebreak'
            })
        elif tag == 'p':
            # End of paragraph - add double line break
            self.text_parts.append({
                'text': '\n\n',
                'formatting': set(),
                'type': 'paragraph_break'
            })
            
    def handle_data(self, data):
        if data.strip():  # Only add non-empty text
            self.text_parts.append({
                'text': data,
                'formatting': self.current_formatting.copy(),
                'type': 'text'
            })
    
    def get_parsed_content(self):
        return self.text_parts


class ExportRequest(BaseModel):
    title: str = Field(..., max_length=MAX_TITLE_LENGTH)
    content: str = Field(..., max_length=MAX_CONTENT_LENGTH)
    format: Literal["txt", "docx", "pdf"]


def sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe file download."""
    # Remove or replace invalid characters
    sanitized = re.sub(r'[<>:"/\\|?*]', '-', filename)
    # Remove multiple consecutive dashes/spaces
    sanitized = re.sub(r'[-\s]+', '-', sanitized)
    # Remove leading/trailing dashes and spaces
    sanitized = sanitized.strip('- ')
    # Ensure filename is not empty
    if not sanitized:
        sanitized = "document"
    # Limit length
    if len(sanitized) > 100:
        sanitized = sanitized[:100]
    return sanitized


def html_to_plain_text(html_content: str) -> str:
    """Convert HTML content from TipTap editor to plain text with paragraph breaks."""
    if not html_content:
        return ""
    
    # Replace paragraph tags with double line breaks
    text = re.sub(r'<p[^>]*>', '', html_content)
    text = re.sub(r'</p>', '\n\n', text)
    
    # Replace br tags with single line breaks
    text = re.sub(r'<br\s*/?>', '\n', text)
    
    # Remove any other HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # Decode HTML entities
    text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
    text = text.replace('&quot;', '"').replace('&#39;', "'")
    
    # Clean up multiple line breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def html_to_rich_text(html_content: str) -> str:
    """Convert HTML content from TipTap editor to rich text with preserved formatting."""
    if not html_content:
        return ""
    
    # Process the HTML to preserve formatting
    text = html_content
    
    # Convert bold tags to markdown-style formatting for intermediate processing
    text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', text, flags=re.DOTALL)
    text = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', text, flags=re.DOTALL)
    
    # Convert italic tags
    text = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', text, flags=re.DOTALL)
    text = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', text, flags=re.DOTALL)
    
    # Convert underline tags (keep as markers for now)
    text = re.sub(r'<u[^>]*>(.*?)</u>', r'_UNDERLINE_START_\1_UNDERLINE_END_', text, flags=re.DOTALL)
    
    # Convert strikethrough
    text = re.sub(r'<s[^>]*>(.*?)</s>', r'~~\1~~', text, flags=re.DOTALL)
    text = re.sub(r'<strike[^>]*>(.*?)</strike>', r'~~\1~~', text, flags=re.DOTALL)
    
    # Handle headings
    for i in range(1, 7):
        text = re.sub(rf'<h{i}[^>]*>(.*?)</h{i}>', rf'{"#" * i} \1\n', text, flags=re.DOTALL)
    
    # Handle lists
    text = re.sub(r'<ul[^>]*>', '\n', text)
    text = re.sub(r'</ul>', '\n', text)
    text = re.sub(r'<ol[^>]*>', '\n', text)
    text = re.sub(r'</ol>', '\n', text)
    text = re.sub(r'<li[^>]*>', '• ', text)
    text = re.sub(r'</li>', '\n', text)
    
    # Handle blockquotes
    text = re.sub(r'<blockquote[^>]*>(.*?)</blockquote>', r'> \1\n', text, flags=re.DOTALL)
    
    # Handle links
    text = re.sub(r'<a[^>]*href=["\']([^"\']*)["\'][^>]*>(.*?)</a>', r'\2 (\1)', text, flags=re.DOTALL)
    
    # Replace paragraph tags with double line breaks
    text = re.sub(r'<p[^>]*>', '', text)
    text = re.sub(r'</p>', '\n\n', text)
    
    # Replace br tags with single line breaks
    text = re.sub(r'<br\s*/?>', '\n', text)
    
    # Remove any remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # Decode HTML entities
    text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
    text = text.replace('&quot;', '"').replace('&#39;', "'").replace('&nbsp;', ' ')
    
    # Clean up multiple line breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def extract_paragraphs(html_content: str) -> list[str]:
    """Extract paragraphs from HTML content for structured formats."""
    if not html_content:
        return [""]
    
    # Find all paragraph tags
    paragraphs = re.findall(r'<p[^>]*>(.*?)</p>', html_content, re.DOTALL)
    
    if not paragraphs:
        # Fallback: treat entire content as one paragraph
        return [html_to_plain_text(html_content)]
    
    processed_paragraphs = []
    for p in paragraphs:
        # Replace br tags with line breaks
        text = re.sub(r'<br\s*/?>', '\n', p)
        # Remove other HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        # Decode HTML entities
        text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
        text = text.replace('&quot;', '"').replace('&#39;', "'").replace('&nbsp;', ' ')
        # Clean and add if not empty
        text = text.strip()
        if text:
            processed_paragraphs.append(text)
    
    return processed_paragraphs if processed_paragraphs else [""]


def extract_rich_paragraphs(html_content: str) -> list[dict]:
    """Extract paragraphs with rich formatting from HTML content."""
    if not html_content:
        return [{"text": "", "formatting": []}]
    
    # Find all paragraph tags and their content
    paragraphs = re.findall(r'<p[^>]*>(.*?)</p>', html_content, re.DOTALL)
    
    if not paragraphs:
        # Fallback: treat entire content as one paragraph
        return [{"text": html_to_rich_text(html_content), "formatting": []}]
    
    processed_paragraphs = []
    for p in paragraphs:
        # Extract formatting information
        formatting_info = []
        
        # Find bold text
        bold_matches = re.finditer(r'<(strong|b)[^>]*>(.*?)</\1>', p, re.DOTALL)
        for match in bold_matches:
            formatting_info.append({"type": "bold", "text": match.group(2)})
        
        # Find italic text
        italic_matches = re.finditer(r'<(em|i)[^>]*>(.*?)</\1>', p, re.DOTALL)
        for match in italic_matches:
            formatting_info.append({"type": "italic", "text": match.group(2)})
        
        # Find underlined text
        underline_matches = re.finditer(r'<u[^>]*>(.*?)</u>', p, re.DOTALL)
        for match in underline_matches:
            formatting_info.append({"type": "underline", "text": match.group(1)})
        
        # Process the paragraph text
        text = p
        # Replace br tags with line breaks
        text = re.sub(r'<br\s*/?>', '\n', text)
        # Keep formatting tags for now, we'll process them in the generators
        
        if text.strip():
            processed_paragraphs.append({
                "text": text,
                "formatting": formatting_info
            })
    
    return processed_paragraphs if processed_paragraphs else [{"text": "", "formatting": []}]


def generate_txt_file(title: str, content: str) -> io.BytesIO:
    """Generate a plain text file with basic formatting preserved."""
    # Parse HTML content
    parser = RichTextParser()
    parser.feed(content)
    text_parts = parser.get_parsed_content()
    
    # Convert to text with basic formatting
    result_text = ""
    for part in text_parts:
        text = part['text']
        formatting = part['formatting']
        
        if 'heading' in formatting:
            # Add markdown-style heading
            level = 1
            for fmt in formatting:
                if fmt.startswith('heading-'):
                    level = int(fmt.split('-')[1])
                    break
            result_text += '#' * level + ' ' + text.strip()
        elif 'bold' in formatting:
            result_text += f"**{text}**"
        elif 'italic' in formatting:
            result_text += f"*{text}*"
        elif 'underline' in formatting:
            result_text += f"_{text}_"
        elif 'strikethrough' in formatting:
            result_text += f"~~{text}~~"
        else:
            result_text += text
    
    # Create the text content
    full_text = f"{title}\n{'=' * len(title)}\n\n{result_text.strip()}"
    
    # Convert to bytes
    file_stream = io.BytesIO()
    file_stream.write(full_text.encode('utf-8'))
    file_stream.seek(0)
    
    return file_stream


def generate_docx_file(title: str, content: str) -> io.BytesIO:
    """Generate a DOCX file with rich formatting preserved."""
    if not docx:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX generation not available"
        )
    
    try:
        # Create a new document
        document = docx.Document()
        
        # Add title
        title_paragraph = document.add_heading(title, level=1)
        title_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a line break
        document.add_paragraph()
        
        # Parse HTML content
        parser = RichTextParser()
        parser.feed(content)
        text_parts = parser.get_parsed_content()
        
        # Build document with formatting
        current_paragraph = document.add_paragraph()
        
        for part in text_parts:
            text = part['text']
            formatting = part['formatting']
            part_type = part['type']
            
            if part_type == 'paragraph_break':
                # Start a new paragraph
                current_paragraph = document.add_paragraph()
            elif part_type == 'linebreak':
                # Add line break within current paragraph
                if current_paragraph.runs:
                    current_paragraph.add_run().add_break()
            elif part_type == 'text' and text.strip():
                # Add formatted text
                run = current_paragraph.add_run(text)
                
                if 'bold' in formatting:
                    run.bold = True
                if 'italic' in formatting:
                    run.italic = True
                if 'underline' in formatting:
                    run.underline = True
                
                # Handle headings by creating a new heading paragraph
                if 'heading' in formatting:
                    # Remove the text from current paragraph
                    current_paragraph._element.remove(run._element)
                    
                    # Determine heading level
                    level = 1
                    for fmt in formatting:
                        if fmt.startswith('heading-'):
                            level = int(fmt.split('-')[1])
                            break
                    
                    # Create heading paragraph
                    heading_para = document.add_heading(text.strip(), level=level)
                    current_paragraph = document.add_paragraph()
        
        # Save to BytesIO
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
        
    except Exception as e:
        raise ValueError(f"Failed to generate DOCX file: {str(e)}")


def generate_pdf_file(title: str, content: str) -> io.BytesIO:
    """Generate a PDF file with rich formatting preserved."""
    if not SimpleDocTemplate:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF generation not available"
        )
    
    try:
        # Create a BytesIO buffer
        file_stream = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            file_stream,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            alignment=1,  # Center alignment
            spaceBefore=0,
            spaceAfter=20
        )
        
        # Create heading styles
        heading_styles = {}
        for i in range(1, 7):
            heading_styles[i] = ParagraphStyle(
                f'Heading{i}',
                parent=styles['Heading1'],
                fontSize=18 - (i-1) * 2,
                spaceBefore=12,
                spaceAfter=6,
                keepWithNext=True
            )
        
        # Build content
        story = []
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Parse HTML content
        parser = RichTextParser()
        parser.feed(content)
        text_parts = parser.get_parsed_content()
        
        # Build paragraphs with formatting
        current_paragraph_text = ""
        
        for part in text_parts:
            text = part['text']
            formatting = part['formatting']
            part_type = part['type']
            
            if part_type == 'paragraph_break':
                # Finish current paragraph and start new one
                if current_paragraph_text.strip():
                    try:
                        story.append(Paragraph(current_paragraph_text, styles['Normal']))
                        story.append(Spacer(1, 0.1*inch))
                    except:
                        # Fallback for problematic formatting
                        clean_text = re.sub(r'<[^>]+>', '', current_paragraph_text)
                        story.append(Paragraph(clean_text, styles['Normal']))
                        story.append(Spacer(1, 0.1*inch))
                current_paragraph_text = ""
            elif part_type == 'linebreak':
                current_paragraph_text += '<br/>'
            elif part_type == 'text' and text.strip():
                # Handle headings specially
                if 'heading' in formatting:
                    # Finish current paragraph first
                    if current_paragraph_text.strip():
                        try:
                            story.append(Paragraph(current_paragraph_text, styles['Normal']))
                        except:
                            clean_text = re.sub(r'<[^>]+>', '', current_paragraph_text)
                            story.append(Paragraph(clean_text, styles['Normal']))
                        current_paragraph_text = ""
                    
                    # Add heading
                    level = 1
                    for fmt in formatting:
                        if fmt.startswith('heading-'):
                            level = int(fmt.split('-')[1])
                            break
                    
                    heading_style = heading_styles.get(level, heading_styles[1])
                    story.append(Paragraph(text.strip(), heading_style))
                    story.append(Spacer(1, 0.1*inch))
                else:
                    # Add formatted text to current paragraph
                    formatted_text = text
                    
                    # Apply formatting tags that ReportLab understands
                    if 'bold' in formatting:
                        formatted_text = f'<b>{formatted_text}</b>'
                    if 'italic' in formatting:
                        formatted_text = f'<i>{formatted_text}</i>'
                    if 'underline' in formatting:
                        formatted_text = f'<u>{formatted_text}</u>'
                    
                    current_paragraph_text += formatted_text
        
        # Add any remaining paragraph
        if current_paragraph_text.strip():
            try:
                story.append(Paragraph(current_paragraph_text, styles['Normal']))
            except:
                clean_text = re.sub(r'<[^>]+>', '', current_paragraph_text)
                story.append(Paragraph(clean_text, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        file_stream.seek(0)
        
        return file_stream
        
    except Exception as e:
        raise ValueError(f"Failed to generate PDF file: {str(e)}")


@router.post("/file")
async def export_file(
    export_request: ExportRequest,
    current_user_id: UUID = Depends(get_current_user)
):
    """
    Export a document to the specified format (txt, docx, pdf).
    Returns the file as a download.
    """
    
    try:
        # Validate content length
        if len(export_request.content) > MAX_CONTENT_LENGTH:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"Content too large. Maximum size: {MAX_CONTENT_LENGTH // 1024}KB"
            )
        
        # Generate filename
        safe_title = sanitize_filename(export_request.title)
        filename = f"{safe_title}.{export_request.format}"
        
        # Generate file based on format
        with sentry_sdk.start_span(
            op="file.export",
            name=f"Export {export_request.format.upper()} file"
        ) as span:
            span.set_data("file.format", export_request.format)
            span.set_data("content.length", len(export_request.content))
            
            if export_request.format == "txt":
                file_stream = generate_txt_file(export_request.title, export_request.content)
                media_type = "text/plain"
            elif export_request.format == "docx":
                file_stream = generate_docx_file(export_request.title, export_request.content)
                media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif export_request.format == "pdf":
                file_stream = generate_pdf_file(export_request.title, export_request.content)
                media_type = "application/pdf"
            else:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Invalid format. Supported formats: txt, docx, pdf"
                )
            
            span.set_data("file.size", file_stream.getbuffer().nbytes)
        
        # Log successful export
        sentry_sdk.add_breadcrumb(
            message=f"Document exported successfully: {filename}",
            category="export",
            level="info",
            data={
                "user_id": str(current_user_id),
                "file_type": export_request.format,
                "file_size": file_stream.getbuffer().nbytes
            }
        )
        
        # Return file as streaming response
        return StreamingResponse(
            io.BytesIO(file_stream.getvalue()),
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename=\"{filename}\""
            }
        )
        
    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        sentry_sdk.capture_exception(e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to export document"
        ) 