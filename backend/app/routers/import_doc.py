# /backend/app/routers/import_doc.py

from typing import Any
from uuid import UUID
import io
import os
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
import sentry_sdk

from ..database import get_db_session
from ..auth import get_current_user
from ..models import Document
from ..schemas import DocumentResponse

# Import parsing libraries
try:
    import docx
except ImportError:
    docx = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

# File size limit (30MB in bytes)
MAX_FILE_SIZE = 30 * 1024 * 1024

# Supported file extensions
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}

router = APIRouter(prefix="/import", tags=["import"])


def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from a plain text file."""
    try:
        text = file_content.decode('utf-8')
    except UnicodeDecodeError:
        # Try other common encodings
        for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
            try:
                text = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Unable to decode text file with supported encodings")
    
    # Convert text to HTML paragraphs for TipTap editor
    # Split by double line breaks first (paragraph separators)
    paragraphs = text.split('\n\n')
    
    # For each paragraph, replace single line breaks with <br> tags
    html_paragraphs = []
    for paragraph in paragraphs:
        if paragraph.strip():
            # Replace single line breaks with <br> tags within paragraphs
            formatted_paragraph = paragraph.strip().replace('\n', '<br>')
            html_paragraphs.append(f'<p>{formatted_paragraph}</p>')
    
    return ''.join(html_paragraphs) if html_paragraphs else '<p></p>'


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from a DOCX file."""
    if not docx:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX parsing not available"
        )
    
    try:
        # Create a BytesIO object from the file content
        file_stream = io.BytesIO(file_content)
        doc = docx.Document(file_stream)
        
        # Extract text from all paragraphs and convert to HTML
        html_paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Handle line breaks within the paragraph text
                formatted_text = text.replace('\n', '<br>')
                html_paragraphs.append(f'<p>{formatted_text}</p>')
        
        return ''.join(html_paragraphs) if html_paragraphs else '<p></p>'
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from a PDF file."""
    if not fitz:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF parsing not available"
        )
    
    try:
        # Open PDF document directly from bytes
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        
        # Extract text from all pages and combine
        all_text = []
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            text = page.get_text()
            if text.strip():
                all_text.append(text.strip())
        
        pdf_document.close()
        
        # Combine all page text
        full_text = '\n\n'.join(all_text)
        
        # More intelligent paragraph detection for PDFs
        # First try splitting by double line breaks
        paragraphs = full_text.split('\n\n')
        
        # If we only get one paragraph (common with PDFs), try different strategies
        if len(paragraphs) == 1:
            lines = full_text.split('\n')
            paragraphs = []
            current_paragraph = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    # Empty line - end current paragraph if it has content
                    if current_paragraph:
                        paragraphs.append(' '.join(current_paragraph))
                        current_paragraph = []
                elif (line.endswith('.') or line.endswith('!') or line.endswith('?')) and len(line) > 50:
                    # Likely end of a sentence/paragraph (long line ending with punctuation)
                    current_paragraph.append(line)
                    paragraphs.append(' '.join(current_paragraph))
                    current_paragraph = []
                elif line[0].isupper() and len(current_paragraph) > 0 and len(' '.join(current_paragraph)) > 100:
                    # New line starting with capital letter, and current paragraph is substantial
                    # Treat as new paragraph
                    paragraphs.append(' '.join(current_paragraph))
                    current_paragraph = [line]
                else:
                    current_paragraph.append(line)
            
            # Add any remaining content
            if current_paragraph:
                paragraphs.append(' '.join(current_paragraph))
        
        html_paragraphs = []
        for paragraph in paragraphs:
            text = paragraph.strip()
            if text:
                # Don't add <br> tags since we've already combined lines into paragraphs
                html_paragraphs.append(f'<p>{text}</p>')
        
        return ''.join(html_paragraphs) if html_paragraphs else '<p></p>'
    except Exception as e:
        raise ValueError(f"Failed to parse PDF file: {str(e)}")


def get_file_extension(filename: str) -> str:
    """Get the file extension from filename."""
    return Path(filename).suffix.lower()


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from uploaded file based on its extension."""
    extension = get_file_extension(filename)
    
    if extension == ".txt":
        return extract_text_from_txt(file_content)
    elif extension == ".docx":
        return extract_text_from_docx(file_content)
    elif extension == ".pdf":
        return extract_text_from_pdf(file_content)
    else:
        raise ValueError(f"Unsupported file extension: {extension}")


@router.post("/file", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def import_file(
    file: UploadFile = File(...),
    current_user_id: UUID = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_session)
):
    """
    Import a document from an uploaded file (.txt, .pdf, .docx).
    Creates a new document with the extracted text content.
    """
    
    # Validate file is provided
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No file provided"
        )
    
    # Check file extension
    extension = get_file_extension(file.filename)
    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file format. Supported formats: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    
    # Read file content
    try:
        file_content = await file.read()
    except Exception as e:
        sentry_sdk.capture_exception(e)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to read uploaded file"
        )
    
    # Check file size
    if len(file_content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File size exceeds maximum limit of {MAX_FILE_SIZE // (1024 * 1024)}MB"
        )
    
    # Extract text from file
    try:
        print(f"Starting to parse {extension} file: {file.filename}")
        print(f"File size: {len(file_content)} bytes")
        
        with sentry_sdk.start_span(
            op="file.parse",
            name=f"Parse {extension} file"
        ) as span:
            span.set_data("file.extension", extension)
            span.set_data("file.size", len(file_content))
            
            extracted_text = extract_text_from_file(file_content, file.filename)
            
            span.set_data("extracted.length", len(extracted_text))
            print(f"Successfully extracted {len(extracted_text)} characters")
            
    except ValueError as e:
        # Log the detailed error for debugging
        print(f"ValueError during file parsing: {str(e)}")
        sentry_sdk.capture_exception(e)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        # Log the detailed error for debugging
        print(f"Unexpected error during file parsing: {str(e)}")
        sentry_sdk.capture_exception(e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process uploaded file: {str(e)}"
        )
    
    # Validate extracted text
    if not extracted_text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No text content found in the uploaded file"
        )
    
    # Create document title from filename (remove extension)
    document_title = Path(file.filename).stem
    
    # Create new document with extracted content
    try:
        new_document = Document(
            profile_id=current_user_id,
            title=document_title,
            content=extracted_text
        )
        
        db.add(new_document)
        await db.commit()
        await db.refresh(new_document)
        
        # Log successful import
        sentry_sdk.add_breadcrumb(
            message=f"Document imported successfully: {document_title}",
            category="import",
            level="info",
            data={
                "document_id": str(new_document.id),
                "file_type": extension,
                "content_length": len(extracted_text)
            }
        )
        
        return DocumentResponse.model_validate(new_document)
        
    except Exception as e:
        sentry_sdk.capture_exception(e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to create document"
        ) 